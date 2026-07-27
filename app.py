"""
Flask web frontend for the Autonomous Research Agent System.
"""
from dotenv import load_dotenv
load_dotenv()

from flask import Flask, render_template, request, jsonify, send_file, abort
from flask_talisman import Talisman
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
import asyncio
import json
import os
import re
import sys
import tempfile
from datetime import datetime
from pathlib import Path
import threading
import concurrent.futures
from werkzeug.utils import secure_filename

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

from src.research_system import AutonomousResearchSystem
from src.agents.custom_citation_formatter import CustomCitationFormatter
from src.agents.literature_builder_agent import LiteratureBuilderAgent
from src.agents.humanizer_agent import HumanizerAgent
from src.agents.latex_citation_reorder import reorder_citations
from src.agents.plagiarism_agent import PlagiarismCheckerAgent
from src.agents.journal_recommender_agent import JournalRecommenderAgent
from src.agents.abstract_screener_agent import AbstractScreenerAgent
from src.agents.checklist_agent import ChecklistAgent
from src.agents.grant_writer_agent import GrantWriterAgent, FUNDER_TEMPLATES
from src.agents.citation_network_agent import CitationNetworkAgent

app = Flask(__name__)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', os.urandom(32).hex())
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'

# ── Rate limiting ────────────────────────────────────────────────────────────
limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["200 per hour", "30 per minute"],
    storage_uri="memory://",
)

# ── Security headers (CSP allows CDN fonts/icons + inline styles needed by app)
csp = {
    'default-src': ["'self'"],
    'script-src':  ["'self'", "'unsafe-inline'",
                    "https://cdn.jsdelivr.net", "https://cdnjs.cloudflare.com",
                    "https://d3js.org"],
    'style-src':   ["'self'", "'unsafe-inline'",
                    "https://cdn.jsdelivr.net", "https://cdnjs.cloudflare.com",
                    "https://fonts.googleapis.com"],
    'font-src':    ["'self'", "https://fonts.gstatic.com",
                    "https://cdnjs.cloudflare.com"],
    'img-src':     ["'self'", "data:", "https:"],
    'connect-src': ["'self'"],
}
Talisman(
    app,
    content_security_policy=csp,
    force_https=False,               # localhost dev — flip to True in production
    strict_transport_security=False, # same
    referrer_policy='strict-origin-when-cross-origin',
    x_content_type_options=True,
    x_xss_protection=True,
)

# ── Input validation helpers ─────────────────────────────────────────────────
_URL_RE = re.compile(r'^https?://', re.I)

def _safe_text(value, max_len: int = 20000) -> str:
    """Strip and truncate text input; reject if suspiciously long."""
    if not isinstance(value, str):
        return ""
    return value.strip()[:max_len]

def _validate_url(url: str) -> str:
    """Return url only if it's a plain http/https web URL, else empty string."""
    url = (url or "").strip()
    if _URL_RE.match(url) and len(url) < 2048:
        return url
    return ""

# Global system instance
research_system = None
executor = concurrent.futures.ThreadPoolExecutor(max_workers=2)

def get_system():
    """Get or create research system instance."""
    global research_system
    if research_system is None:
        research_system = AutonomousResearchSystem()
    return research_system


def run_research_in_thread(topic):
    """Run research in a separate thread with its own event loop."""
    def research_task():
        # Create new event loop for this thread
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        try:
            system = get_system()
            result = loop.run_until_complete(system.research(topic))
            return result
        except Exception as e:
            raise e
        finally:
            loop.close()
    
    return research_task()


def _generate_custom_citation(paper):
    """Generate custom citation format with accurate bibliographic data."""
    # Generate bibitem key
    bibitem_key = _generate_bibitem_key(paper)
    
    # Format authors
    formatted_authors = _format_authors_custom(paper.authors)
    
    # Format title
    formatted_title = _format_title_custom(paper.title)
    
    # Get accurate volume, issue, and page information
    volume_info = _extract_volume_info_crossref(paper)
    
    # Get paper URL
    paper_url = _get_paper_url(paper)
    
    # Create citation with accurate data
    citation = f"\\bibitem{{{bibitem_key}}} {formatted_authors}, {formatted_title}, {paper.venue}"
    
    # Add volume and issue information ONLY if we have real data
    if volume_info['volume']:
        citation += f" \\textbf{{{volume_info['volume']}}}"
        if volume_info['issue']:
            citation += f"({volume_info['issue']})"
    
    # Add year
    citation += f" ({paper.year})"
    
    # Add pages ONLY if we have real data
    if volume_info['pages']:
        citation += f" {volume_info['pages']}"
    
    if paper_url:
        citation += f". {paper_url}"
    
    return citation


def _extract_volume_info_crossref(paper):
    """Extract accurate volume, issue, and page information using CrossRef API."""
    volume_info = {
        'volume': None,
        'issue': None,
        'pages': None,
        'article_number': None
    }
    
    # Try to extract from DOI using CrossRef API
    if paper.doi:
        try:
            import requests
            import time
            
            # Clean DOI - handle various DOI formats
            doi = paper.doi.strip()
            if doi.startswith('https://doi.org/'):
                doi = doi.replace('https://doi.org/', '')
            elif doi.startswith('http://dx.doi.org/'):
                doi = doi.replace('http://dx.doi.org/', '')
            elif doi.startswith('doi:'):
                doi = doi.replace('doi:', '')
            
            # Query CrossRef API for accurate bibliographic data
            crossref_url = f"https://api.crossref.org/works/{doi}"
            headers = {
                'User-Agent': 'Research System/1.0 (mailto:research@example.com)',
                'Accept': 'application/json'
            }
            
            print(f"Fetching bibliographic data for DOI: {doi}")
            response = requests.get(crossref_url, headers=headers, timeout=10)
            
            if response.status_code == 200:
                data = response.json()
                work = data.get('message', {})
                
                # Extract volume
                if 'volume' in work and work['volume']:
                    volume_info['volume'] = str(work['volume']).strip()
                
                # Extract issue
                if 'issue' in work and work['issue']:
                    volume_info['issue'] = str(work['issue']).strip()
                
                # Extract pages - handle different formats
                if 'page' in work and work['page']:
                    pages = str(work['page']).strip()
                    if pages:
                        volume_info['pages'] = pages
                elif 'article-number' in work and work['article-number']:
                    article_num = str(work['article-number']).strip()
                    volume_info['article_number'] = article_num
                    volume_info['pages'] = article_num
                
                print(f"Successfully extracted: vol={volume_info['volume']}, issue={volume_info['issue']}, pages={volume_info['pages']}")
                
                # Small delay to be respectful to CrossRef API
                time.sleep(0.1)
                
            elif response.status_code == 404:
                print(f"DOI {doi} not found in CrossRef (404)")
            else:
                print(f"CrossRef API returned status {response.status_code} for DOI {doi}")
                
        except Exception as e:
            print(f"Error fetching bibliographic data for DOI {paper.doi}: {str(e)}")
    
    # Try to extract from venue name if it contains volume/issue info
    if not any([volume_info['volume'], volume_info['issue'], volume_info['pages']]):
        if paper.venue:
            venue_lower = paper.venue.lower()
            import re
            
            # Look for patterns like "vol. 25, no. 3" or "25(3)" in venue
            vol_match = re.search(r'vol(?:ume)?\.?\s*(\d+)', venue_lower)
            if vol_match:
                volume_info['volume'] = vol_match.group(1)
            
            issue_match = re.search(r'(?:no|issue)\.?\s*(\d+)', venue_lower)
            if issue_match:
                volume_info['issue'] = issue_match.group(1)
            
            # Look for pattern like "25(3)" 
            vol_issue_match = re.search(r'(\d+)\((\d+)\)', venue_lower)
            if vol_issue_match:
                volume_info['volume'] = vol_issue_match.group(1)
                volume_info['issue'] = vol_issue_match.group(2)
    
    return volume_info


def _generate_bibitem_key(paper):
    """Generate bibitem key from authors and year."""
    if not paper.authors:
        return f"Unknown{paper.year % 100:02d}"
    
    # Get first 2-3 authors
    authors_for_key = paper.authors[:3] if len(paper.authors) >= 3 else paper.authors[:2]
    
    # Extract first two letters of last names
    key_parts = []
    for author in authors_for_key:
        last_name = author.split()[-1] if author.split() else author
        if len(last_name) >= 2:
            key_parts.append(last_name[:2].title())
        elif len(last_name) == 1:
            key_parts.append(last_name.upper())
    
    # Add year (last 2 digits)
    year_suffix = paper.year % 100
    
    key = "".join(key_parts) + f"{year_suffix:02d}"
    return key


def _format_authors_custom(authors):
    """Format authors with initials first, then last name."""
    if not authors:
        return "Unknown Author"
    
    formatted_authors = []
    
    for author in authors:
        author = author.strip()
        parts = author.split()
        
        if len(parts) < 2:
            formatted_authors.append(author)
            continue
        
        # Check if author is already in correct format (e.g., "Y. Saad", "M.H. Schultz")
        if len(parts) == 2 and len(parts[0]) <= 4 and '.' in parts[0]:
            # Already in format like "Y. Saad" or "M.H. Schultz"
            formatted_authors.append(author)
            continue
            
        # Last part is last name
        last_name = parts[-1]
        
        # Everything else is first/middle names
        first_middle = parts[:-1]
        
        # Create initials
        initials = []
        for name in first_middle:
            if name and name[0].isalpha():
                # Handle names that might already have periods
                if name.endswith('.'):
                    initials.append(name)
                else:
                    initials.append(f"{name[0].upper()}.")
        
        # Format as "F.M. Last" (initials first, then last name)
        if initials:
            formatted_authors.append(f"{''.join(initials)} {last_name}")
        else:
            formatted_authors.append(last_name)
    
    return ", ".join(formatted_authors)


def _format_title_custom(title):
    """Format title with only first letter capital, except proper nouns."""
    if not title:
        return ""
    
    # Clean title
    title = title.strip()
    if title.endswith('.'):
        title = title[:-1]
    
    # Convert to lowercase first
    title_lower = title.lower()
    
    # Capitalize first letter
    if title_lower:
        title_formatted = title_lower[0].upper() + title_lower[1:]
    else:
        title_formatted = title_lower
    
    # Capitalize proper nouns
    proper_nouns = ["wiener", "euler", "hamilton", "fibonacci", "pascal", "newton", "gauss",
                   "fourier", "laplace", "bayes", "markov", "poisson", "bernoulli",
                   "covid", "sars", "hiv", "aids", "dna", "rna", "pcr", "crispr"]
    
    for proper_noun in proper_nouns:
        import re
        pattern = r'\b' + re.escape(proper_noun.lower()) + r'\b'
        replacement = proper_noun.capitalize()
        title_formatted = re.sub(pattern, replacement, title_formatted, flags=re.IGNORECASE)
    
    return title_formatted


def _get_paper_url(paper):
    """Get the best URL for the paper."""
    if paper.doi:
        if paper.doi.startswith('http'):
            return paper.doi
        else:
            return f"https://doi.org/{paper.doi}"
    elif paper.url:
        return paper.url
    elif paper.arxiv_id:
        return f"https://arxiv.org/abs/{paper.arxiv_id}"
    else:
        return ""


@app.route('/')
def index():
    """Main page."""
    return render_template('index.html')


@app.route('/research', methods=['POST'])
def research():
    """Perform research on a topic."""
    try:
        data = request.get_json()
        topic = data.get('topic', '').strip()
        
        if not topic:
            return jsonify({'error': 'Topic is required'}), 400
        
        print(f"Starting research on topic: {topic}")
        
        # Run research in a separate thread to avoid event loop conflicts
        future = executor.submit(run_research_in_thread, topic)
        results = future.result(timeout=300)  # 5 minute timeout
        
        print(f"Research completed successfully")
        
        # Convert results to JSON-serializable format
        response_data = {
            'topic': results.topic_map.main_topic,
            'generated_at': results.generated_at.isoformat(),
            'summary': {
                'papers_analyzed': results.total_papers_analyzed,
                'claims_extracted': results.total_claims_extracted,
                'contradictions_found': len(results.contradictions),
                'research_gaps_identified': len(results.research_gaps)
            },
            'usage': results.usage,
            'run_metadata': results.run_metadata,
            'topic_map': {
                'main_topic': results.topic_map.main_topic,
                'subtopics': results.topic_map.subtopics,
                'methods': results.topic_map.methods,
                'keywords': results.topic_map.keywords,
                'datasets': results.topic_map.datasets
            },
            'papers': [
                {
                    'title': paper.title,
                    'authors': paper.authors,
                    'year': paper.year,
                    'venue': paper.venue,
                    'relevance_score': paper.relevance_score,
                    'abstract': paper.abstract[:300] + '...' if len(paper.abstract) > 300 else paper.abstract,
                    'doi': paper.doi,
                    'arxiv_id': paper.arxiv_id,
                    'url': paper.url
                }
                for paper in results.papers
            ],
            'claims': [
                {
                    'statement': claim.statement,
                    'confidence': claim.confidence,
                    'metrics': claim.metrics,
                    'datasets': claim.datasets,
                    'conditions': claim.conditions
                }
                for claim in results.claims
            ],
            'contradictions': [
                {
                    'explanation': contradiction.explanation,
                    'type': contradiction.contradiction_type,
                    'severity': contradiction.severity,
                    'likert_score': contradiction.likert_score,
                    'verdict': contradiction.verdict,
                    'evidence_claim1': contradiction.evidence_claim1,
                    'evidence_claim2': contradiction.evidence_claim2,
                    'paper1_id': contradiction.paper1_id,
                    'paper2_id': contradiction.paper2_id,
                    'detection_method': contradiction.detection_method
                }
                for contradiction in results.contradictions
            ],
            'research_gaps': [
                {
                    'description': gap.description,
                    'type': gap.gap_type,
                    'priority': gap.priority,
                    'potential_questions': gap.potential_questions,
                    'evidence': gap.evidence,
                    'rank': gap.rank,
                    'detection_method': gap.detection_method
                }
                for gap in results.research_gaps
            ],
            'citations': [
                {
                    'paper_id': citation.paper_id,
                    'bibtex': citation.bibtex,
                    'apa': citation.apa,
                    'ieee': citation.ieee,
                    'mla': citation.mla if hasattr(citation, 'mla') else None,
                    'custom_format': _generate_custom_citation(paper),
                    'paper_url': _get_paper_url(paper),
                    'journal_name': paper.venue,
                    'year': paper.year
                }
                for citation, paper in zip(results.citations, results.papers)
            ]
        }
        
        # Save results to file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"research_{timestamp}.json"
        filepath = Path("output") / filename
        
        with open(filepath, 'w') as f:
            json.dump(response_data, f, indent=2, default=str)
        
        response_data['download_url'] = f'/download/{filename}'
        
        return jsonify(response_data)
        
    except concurrent.futures.TimeoutError:
        return jsonify({'error': 'Research timed out. Please try a more specific topic.'}), 500
    except Exception as e:
        print(f"Error in research: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Research failed: {str(e)}'}), 500


@app.route('/download/<filename>')
def download_file(filename):
    """Download research results file."""
    try:
        filepath = Path("output") / filename
        if filepath.exists():
            return send_file(filepath, as_attachment=True)
        else:
            return jsonify({'error': 'File not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/history')
def history():
    """Get list of previous research results."""
    try:
        output_dir = Path("output")
        files = []
        
        if output_dir.exists():
            for file in output_dir.glob("research_*.json"):
                try:
                    with open(file, 'r') as f:
                        data = json.load(f)
                    
                    files.append({
                        'filename': file.name,
                        'topic': data.get('topic', 'Unknown'),
                        'generated_at': data.get('generated_at', ''),
                        'papers_count': len(data.get('papers', [])),
                        'claims_count': len(data.get('claims', [])),
                        'download_url': f'/download/{file.name}'
                    })
                except:
                    continue
        
        # Sort by date (newest first)
        files.sort(key=lambda x: x['generated_at'], reverse=True)
        
        return jsonify(files)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/citations')
def citations_page():
    """Enhanced citations page."""
    return render_template('enhanced_citations.html')


@app.route('/literature')
def literature_page():
    """Literature builder page."""
    return render_template('literature_builder.html')


@app.route('/corrections-details')
def corrections_details():
    """Corrections details page."""
    return render_template('corrections_details.html')


@app.route('/validator')
def reference_validator_page():
    """Reference validator page."""
    return render_template('reference_validator.html')


@app.route('/validate-references', methods=['POST'])
def validate_references():
    """Validate and correct uploaded reference file."""
    try:
        data = request.get_json()
        content = data.get('content', '').strip()
        file_format = data.get('format', 'bibtex')
        options = data.get('options', {})
        
        if not content:
            return jsonify({'error': 'No content provided'}), 400
        
        app.logger.info(f"Reference validation request: format={file_format}, length={len(content)}")
        
        # Run validation in a separate thread
        future = executor.submit(run_reference_validation, content, file_format, options)
        results = future.result(timeout=300)  # 5 minute timeout
        
        app.logger.info(f"Reference validation completed successfully")
        
        return jsonify(results)
        
    except concurrent.futures.TimeoutError:
        app.logger.error(f"Reference validation timed out")
        return jsonify({'error': 'Reference validation timed out. Please try with fewer references.'}), 500
    except Exception as e:
        app.logger.error(f"Error in reference validation: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Reference validation failed: {str(e)}'}), 500


@app.route('/reformat-references', methods=['POST'])
def reformat_references():
    """Reformat references to different output format."""
    try:
        data = request.get_json()
        references_data = data.get('references', [])
        output_format = data.get('format', 'bibitem')
        
        # Import reference validator
        from src.agents.reference_validator import ReferenceValidator
        validator = ReferenceValidator()
        
        # Create a mock result object with the references
        class MockResult:
            def __init__(self, refs):
                self.corrected_references = refs
        
        mock_result = MockResult(references_data)
        formatted_output = validator.generate_corrected_file(mock_result, output_format)
        
        return jsonify({
            'formatted_references': formatted_output,
            'format': output_format
        })
        
    except Exception as e:
        app.logger.error(f"Error reformatting references: {str(e)}")
        return jsonify({'error': f'Reformatting failed: {str(e)}'}), 500


def run_reference_validation(content: str, file_format: str, options: dict):
    """Run reference validation in a separate thread."""
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    
    try:
        # Import reference validator
        from src.agents.reference_validator import ReferenceValidator, ReferenceValidationResult
        
        app.logger.info(f"Starting reference validation")
        validator = ReferenceValidator()
        
        # Process the reference file
        result = loop.run_until_complete(validator.process_reference_file(content, file_format))
        
        app.logger.info(f"Reference validation completed: {result.original_count} → {result.final_count}")
        
        # Generate outputs in different formats
        bibitem_output = validator.generate_corrected_file(result, 'bibitem')
        bibtex_output = validator.generate_corrected_file(result, 'bibtex')
        plain_output = validator.generate_corrected_file(result, 'plain')
        validation_report = validator.generate_validation_report(result)
        
        # Prepare response
        response = {
            'stats': {
                'original_count': result.original_count,
                'duplicates_removed': len(result.duplicates_removed),
                'format_corrections': len(result.format_corrections),
                'spelling_corrections': len(result.spelling_corrections),
                'invalid_papers': len(result.invalid_papers),
                'final_count': result.final_count
            },
            'processing_log': result.processing_log,
            'corrections': [],
            'issues': [],
            'corrected_references': bibitem_output,  # Default format
            'corrected_references_data': result.corrected_references,
            'outputs': {
                'bibitem': bibitem_output,
                'bibtex': bibtex_output,
                'plain': plain_output
            },
            'validation_report': validation_report
        }
        
        # Process corrections for display - avoid duplicates
        corrections_by_ref = {}
        
        # Group format corrections
        for correction in result.format_corrections:
            ref_key = correction['reference_key']
            if ref_key not in corrections_by_ref:
                corrections_by_ref[ref_key] = {
                    'type': 'Format',
                    'reference_key': ref_key,
                    'details': [],
                    'corrections_count': 0,
                    'processed_fields': set()  # Track processed fields to avoid duplicates
                }
            
            for corr_text in correction['corrections']:
                # Parse correction text to extract before/after
                if ' → ' in corr_text:
                    field, change = corr_text.split(':', 1) if ':' in corr_text else ('Field', corr_text)
                    field = field.strip()
                    
                    # Skip if this field was already processed
                    if field in corrections_by_ref[ref_key]['processed_fields']:
                        continue
                        
                    if ' → ' in change:
                        before, after = change.split(' → ', 1)
                        corrections_by_ref[ref_key]['details'].append({
                            'field': field,
                            'before': before.strip().strip("'\""),
                            'after': after.strip().strip("'\"")
                        })
                        corrections_by_ref[ref_key]['corrections_count'] += 1
                        corrections_by_ref[ref_key]['processed_fields'].add(field)
        
        # Group spelling corrections
        for correction in result.spelling_corrections:
            ref_key = correction['reference_key']
            if ref_key not in corrections_by_ref:
                corrections_by_ref[ref_key] = {
                    'type': 'Spelling',
                    'reference_key': ref_key,
                    'details': [],
                    'corrections_count': 0,
                    'processed_fields': set()
                }
            
            for corr_text in correction['corrections']:
                if ' → ' in corr_text:
                    field, change = corr_text.split(':', 1) if ':' in corr_text else ('Spelling', corr_text)
                    field = field.strip()
                    
                    if field in corrections_by_ref[ref_key]['processed_fields']:
                        continue
                        
                    if ' → ' in change:
                        before, after = change.split(' → ', 1)
                        corrections_by_ref[ref_key]['details'].append({
                            'field': field,
                            'before': before.strip().strip("'\""),
                            'after': after.strip().strip("'\"")
                        })
                        corrections_by_ref[ref_key]['corrections_count'] += 1
                        corrections_by_ref[ref_key]['processed_fields'].add(field)
        
        # Group paper verification corrections - only add unique corrections
        for verification in result.verification_results:
            ref_key = verification['reference_key']
            ver_result = verification['verification']
            
            if ver_result.get('corrections_made'):
                if ref_key not in corrections_by_ref:
                    corrections_by_ref[ref_key] = {
                        'type': 'Paper Data',
                        'reference_key': ref_key,
                        'details': [],
                        'corrections_count': 0,
                        'processed_fields': set()
                    }
                elif corrections_by_ref[ref_key]['type'] != 'Paper Data':
                    corrections_by_ref[ref_key]['type'] = 'Format & Data'
                
                for corr_text in ver_result['corrections_made']:
                    if 'corrected:' in corr_text or 'added:' in corr_text:
                        # Parse correction text
                        if 'corrected:' in corr_text:
                            field, change = corr_text.split(' corrected:', 1)
                        elif 'added:' in corr_text:
                            field, change = corr_text.split(' added:', 1)
                        else:
                            continue
                            
                        field = field.strip()
                        
                        # Skip if this field was already processed
                        if field in corrections_by_ref[ref_key]['processed_fields']:
                            continue
                            
                        if ' → ' in change:
                            before, after = change.split(' → ', 1)
                        else:
                            before = 'Not provided'
                            after = change.strip()
                        
                        corrections_by_ref[ref_key]['details'].append({
                            'field': field,
                            'before': before.strip().strip("'\""),
                            'after': after.strip().strip("'\"")
                        })
                        corrections_by_ref[ref_key]['corrections_count'] += 1
                        corrections_by_ref[ref_key]['processed_fields'].add(field)
        
        # Clean up processed_fields from response and convert to list
        for ref_key in corrections_by_ref:
            del corrections_by_ref[ref_key]['processed_fields']
        
        response['corrections'] = list(corrections_by_ref.values())
        
        # Process issues for display
        for duplicate in result.duplicates_removed:
            response['issues'].append({
                'type': 'Duplicate',
                'reference_key': duplicate['reference'].get('key', 'unknown'),
                'description': duplicate['reason']
            })
        
        for invalid in result.invalid_papers:
            response['issues'].append({
                'type': 'Invalid',
                'reference_key': invalid['reference'].get('key', 'unknown'),
                'description': invalid['reason']
            })
        
        app.logger.info(f"Reference validation response prepared successfully")
        return response
        
    except Exception as e:
        app.logger.error(f"Error in reference validation: {str(e)}")
        raise e
        
    finally:
        loop.close()


@app.route('/generate-literature', methods=['POST'])
def generate_literature():
    """Generate structured literature from research results."""
    try:
        data = request.get_json()
        topic = data.get('topic', '').strip()
        filters = data.get('filters', {})
        
        if not topic:
            return jsonify({'error': 'Topic is required'}), 400
        
        app.logger.info(f"Literature generation request for topic: {topic}")
        
        # Run research and literature generation in a separate thread
        future = executor.submit(run_literature_generation, topic, filters)
        results = future.result(timeout=300)  # 5 minute timeout
        
        app.logger.info(f"Literature generated successfully for topic: {topic}")
        
        return jsonify(results)
        
    except concurrent.futures.TimeoutError:
        app.logger.error(f"Literature generation timed out for topic: {topic}")
        return jsonify({'error': 'Literature generation timed out. Please try a more specific topic.'}), 500
    except Exception as e:
        app.logger.error(f"Error in literature generation for topic '{topic}': {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Literature generation failed: {str(e)}'}), 500


def run_literature_generation(topic: str, filters: dict):
    """Run literature generation in a separate thread."""
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    
    try:
        # Get research system and run research
        app.logger.info(f"Starting research for topic: {topic}")
        system = get_system()
        research_results = loop.run_until_complete(system.research(topic))
        
        app.logger.info(f"Research completed. Found {len(research_results.papers)} papers, {len(research_results.claims)} claims")
        
        # Generate literature
        app.logger.info("Starting literature generation...")
        literature_agent = LiteratureBuilderAgent()
        literature_document = loop.run_until_complete(literature_agent.process(research_results))
        
        app.logger.info(f"Literature generation completed. Generated {len(literature_document.sections)} sections")
        
        # Get statistics
        stats = literature_agent.get_literature_stats(literature_document)
        
        # Convert to JSON-serializable format
        result = {
            'topic': topic,
            'outline': {
                'title': literature_document.outline.title,
                'sections': literature_document.outline.sections,
                'total_papers': literature_document.outline.total_papers,
                'total_claims': literature_document.outline.total_claims,
                'date_range': literature_document.outline.date_range,
                'estimated_word_count': getattr(literature_document.outline, 'estimated_word_count', 0)
            },
            'sections': [
                {
                    'section_type': section.section_type,
                    'title': section.title,
                    'content': section.content,
                    'citations': section.citations,
                    'claim_ids': section.claim_ids,
                    'word_count': section.word_count
                }
                for section in literature_document.sections
            ],
            'bibliography': literature_document.bibliography,
            'metadata': literature_document.metadata,
            'stats': stats,
            'generated_at': literature_document.generated_at.isoformat()
        }
        
        app.logger.info(f"Literature generation result prepared successfully")
        return result
        
    except Exception as e:
        app.logger.error(f"Error in literature generation: {str(e)}")
        raise e
        
    finally:
        loop.close()


@app.route('/generate-custom-citations', methods=['POST'])
def generate_custom_citations():
    """Generate custom formatted citations."""
    try:
        data = request.get_json()
        topic = data.get('topic', '').strip()
        
        if not topic:
            return jsonify({'error': 'Topic is required'}), 400
        
        print(f"Generating custom citations for topic: {topic}")
        
        # Run research and citation generation in a separate thread
        future = executor.submit(run_citation_generation, topic)
        results = future.result(timeout=300)  # 5 minute timeout
        
        print(f"Custom citations generated successfully")
        
        return jsonify(results)
        
    except concurrent.futures.TimeoutError:
        return jsonify({'error': 'Citation generation timed out. Please try a more specific topic.'}), 500
    except Exception as e:
        print(f"Error in citation generation: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Citation generation failed: {str(e)}'}), 500


def run_citation_generation(topic):
    """Run citation generation in a separate thread."""
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    
    try:
        # Get research system and run research
        system = get_system()
        results = loop.run_until_complete(system.research(topic))
        
        # Generate custom citations
        formatter = CustomCitationFormatter()
        custom_citations = loop.run_until_complete(formatter.process(results.papers))
        
        # Get statistics
        stats = formatter.get_citation_stats(custom_citations)
        
        return {
            'topic': topic,
            'citations': custom_citations,
            'stats': stats,
            'total_papers': len(results.papers)
        }
        
    finally:
        loop.close()


@app.route('/debug')
def debug_form():
    """Debug form page."""
    return send_file('debug_form.html')


@app.route('/test')
def test():
    """System status — returns JSON consumed by the status popup in the UI."""
    import platform, sys
    from src.agents.llm_client import provider_info, _ollama_available
    try:
        system = get_system()
        llm_info = provider_info()
        llm_available = llm_info["available"]
        llm_label = llm_info["name"] if llm_available else "None"

        apis = {
            "ArXiv": True,
            "OpenAlex": True,
            "CrossRef": True,
            "PubMed": True,
            f"LLM ({llm_label})": llm_available,
            "Ollama (local)": _ollama_available(),
        }
        return jsonify({
            'status': 'ok',
            'message': 'All systems operational',
            'system_initialized': system is not None,
            'python_version': sys.version.split()[0],
            'platform': platform.system(),
            'llm_provider': llm_info,
            'apis': apis,
            'features': {
                'llm_extraction': llm_available,
                'humanizer': llm_available,
                'pdf_export': True,
                'docx_export': True,
                'abstract_screener': llm_available,
                'grant_writer': llm_available,
                'checklist': llm_available,
            }
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': str(e),
            'apis': {},
            'features': {}
        }), 500


@app.route('/humanizer')
def humanizer_page():
    from src.agents.llm_client import provider_info
    llm = provider_info()
    return render_template('humanizer.html', llm_info=llm)


@limiter.limit("10 per minute")
@app.route('/humanize', methods=['POST'])
def humanize():
    try:
        text = None

        if request.content_type and 'multipart/form-data' in request.content_type:
            file = request.files.get('file')
            if not file or file.filename == '':
                return jsonify({'error': 'No file uploaded'}), 400

            filename = secure_filename(file.filename)
            if not filename.lower().endswith('.pdf'):
                return jsonify({'error': 'Only PDF files are accepted'}), 400

            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name

            try:
                import fitz
                doc = fitz.open(tmp_path)
                pages_text = []
                for page in doc:
                    pages_text.append(page.get_text())
                doc.close()
                text = '\n\n'.join(pages_text).strip()
            finally:
                os.unlink(tmp_path)

            if not text:
                return jsonify({'error': 'Could not extract text from the PDF'}), 422
            method_source = 'pdf'
        else:
            data = request.get_json(silent=True) or {}
            text = (data.get('text') or '').strip()
            if not text:
                return jsonify({'error': 'No text provided'}), 400
            method_source = 'text'

        agent = HumanizerAgent()
        result = agent.humanize(text)
        result['source'] = method_source
        return jsonify(result)

    except Exception as e:
        app.logger.error(f'Humanize error: {e}')
        import traceback; traceback.print_exc()
        return jsonify({'error': f'Humanization failed: {str(e)}'}), 500


@app.route('/latex-reorder')
def latex_reorder_page():
    return render_template('latex_reorder.html')


@app.route('/latex-reorder', methods=['POST'])
def latex_reorder():
    try:
        file = request.files.get('file')
        if not file or file.filename == '':
            return jsonify({'error': 'No file uploaded'}), 400

        filename = secure_filename(file.filename)
        source = file.read().decode('utf-8', errors='replace')

        reordered_source, info = reorder_citations(source)

        if 'error' in info:
            return jsonify({'error': info['error']}), 422

        return jsonify({
            'reordered_source': reordered_source,
            'citation_map': info.get('citation_map', {}),
            'total_citations': info.get('total_citations', 0),
            'keys_found_in_body': info.get('keys_found_in_body', 0),
        })

    except UnicodeDecodeError:
        return jsonify({'error': 'File encoding not supported. Please save as UTF-8.'}), 422
    except Exception as e:
        app.logger.error(f'LaTeX reorder error: {e}')
        import traceback; traceback.print_exc()
        return jsonify({'error': f'Reordering failed: {str(e)}'}), 500


@app.route('/export/json', methods=['POST'])
def export_json():
    try:
        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        topic_slug = (data.get('topic') or 'research').replace(' ', '_')[:40]
        filename = f'research_{topic_slug}_{timestamp}.json'

        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as tmp:
            json.dump(data, tmp, indent=2, default=str)
            tmp_path = tmp.name

        return send_file(
            tmp_path,
            mimetype='application/json',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        app.logger.error(f'JSON export error: {e}')
        return jsonify({'error': f'Export failed: {str(e)}'}), 500


@app.route('/export/pdf', methods=['POST'])
def export_pdf():
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        from reportlab.lib.enums import TA_LEFT, TA_CENTER

        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        topic_slug = (data.get('topic') or 'research').replace(' ', '_')[:40]
        filename = f'research_{topic_slug}_{timestamp}.pdf'

        tmp_path = os.path.join(tempfile.gettempdir(), filename)

        doc = SimpleDocTemplate(
            tmp_path,
            pagesize=A4,
            rightMargin=2*cm, leftMargin=2*cm,
            topMargin=2*cm, bottomMargin=2*cm
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle('Title', parent=styles['Title'], fontSize=18, spaceAfter=6)
        h2_style = ParagraphStyle('H2', parent=styles['Heading2'], fontSize=13, spaceBefore=14, spaceAfter=4)
        body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=10, leading=14, spaceAfter=4)
        meta_style = ParagraphStyle('Meta', parent=styles['Normal'], fontSize=9, textColor=colors.HexColor('#666666'), spaceAfter=2)

        story = []
        topic = data.get('topic', 'Research Report')
        generated_at = data.get('generated_at', datetime.now().isoformat())

        story.append(Paragraph(f'Research Report: {topic}', title_style))
        story.append(Paragraph(f'Generated: {generated_at}', meta_style))
        story.append(HRFlowable(width='100%', thickness=1, color=colors.HexColor('#30363d'), spaceAfter=12))

        summary = data.get('summary', {})
        if summary:
            story.append(Paragraph('Summary', h2_style))
            summary_data = [
                ['Papers Analyzed', str(summary.get('papers_analyzed', 0))],
                ['Claims Extracted', str(summary.get('claims_extracted', 0))],
                ['Contradictions Found', str(summary.get('contradictions_found', 0))],
                ['Research Gaps', str(summary.get('research_gaps_identified', 0))],
            ]
            t = Table(summary_data, colWidths=[8*cm, 4*cm])
            t.setStyle(TableStyle([
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.HexColor('#444444')),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.HexColor('#f8f8f8'), colors.white]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
                ('PADDING', (0, 0), (-1, -1), 5),
            ]))
            story.append(t)
            story.append(Spacer(1, 12))

        papers = data.get('papers', [])
        if papers:
            story.append(Paragraph(f'Papers ({len(papers)})', h2_style))
            for i, p in enumerate(papers, 1):
                title = p.get('title', 'Untitled')
                authors = ', '.join(p.get('authors', []))
                year = p.get('year', '')
                venue = p.get('venue', '')
                abstract = p.get('abstract', '')
                story.append(Paragraph(f'{i}. {title}', ParagraphStyle('PaperTitle', parent=body_style, fontName='Helvetica-Bold')))
                story.append(Paragraph(f'{authors} · {year} · {venue}', meta_style))
                if abstract:
                    story.append(Paragraph(abstract[:400] + ('...' if len(abstract) > 400 else ''), body_style))
                story.append(Spacer(1, 6))

        gaps = data.get('research_gaps', [])
        if gaps:
            story.append(Paragraph(f'Research Gaps ({len(gaps)})', h2_style))
            for g in gaps[:10]:
                story.append(Paragraph(f'• {g.get("description", "")}', body_style))

        citations = data.get('citations', [])
        if citations:
            story.append(Paragraph(f'Citations ({len(citations)})', h2_style))
            for c in citations:
                text = c.get('custom_format') or c.get('bibtex') or ''
                if text:
                    story.append(Paragraph(text, ParagraphStyle('Cite', parent=body_style, fontName='Courier', fontSize=8, leading=11)))
                    story.append(Spacer(1, 4))

        doc.build(story)

        return send_file(
            tmp_path,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=filename
        )

    except ImportError:
        return jsonify({'error': 'reportlab is not installed. Run: pip install reportlab'}), 500
    except Exception as e:
        app.logger.error(f'PDF export error: {e}')
        import traceback; traceback.print_exc()
        return jsonify({'error': f'PDF export failed: {str(e)}'}), 500


@app.route('/export/docx', methods=['POST'])
def export_docx():
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        data = request.get_json(silent=True)
        if not data:
            return jsonify({'error': 'No data provided'}), 400

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        topic_slug = (data.get('topic') or 'research').replace(' ', '_')[:40]
        filename = f'research_{topic_slug}_{timestamp}.docx'
        tmp_path = os.path.join(tempfile.gettempdir(), filename)

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3)
            section.right_margin = Cm(3)

        topic = data.get('topic', 'Research Report')
        generated_at = data.get('generated_at', datetime.now().isoformat())

        heading = doc.add_heading(f'Research Report: {topic}', level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        meta = doc.add_paragraph(f'Generated: {generated_at}')
        meta.runs[0].font.size = Pt(9)
        meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

        summary = data.get('summary', {})
        if summary:
            doc.add_heading('Summary', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            rows_data = [
                ('Papers Analyzed', str(summary.get('papers_analyzed', 0))),
                ('Claims Extracted', str(summary.get('claims_extracted', 0))),
                ('Contradictions Found', str(summary.get('contradictions_found', 0))),
                ('Research Gaps', str(summary.get('research_gaps_identified', 0))),
            ]
            for i, (label, value) in enumerate(rows_data):
                row = table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = value
            doc.add_paragraph()

        papers = data.get('papers', [])
        if papers:
            doc.add_heading(f'Papers ({len(papers)})', level=1)
            for i, p in enumerate(papers, 1):
                title_para = doc.add_paragraph()
                run = title_para.add_run(f'{i}. {p.get("title", "Untitled")}')
                run.bold = True
                run.font.size = Pt(11)

                authors = ', '.join(p.get('authors', []))
                year = p.get('year', '')
                venue = p.get('venue', '')
                meta_para = doc.add_paragraph(f'{authors} · {year} · {venue}')
                meta_para.runs[0].font.size = Pt(9)
                meta_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                abstract = p.get('abstract', '')
                if abstract:
                    abs_para = doc.add_paragraph(abstract[:500] + ('...' if len(abstract) > 500 else ''))
                    abs_para.runs[0].font.size = Pt(10)

                doc.add_paragraph()

        gaps = data.get('research_gaps', [])
        if gaps:
            doc.add_heading(f'Research Gaps ({len(gaps)})', level=1)
            for g in gaps[:10]:
                p = doc.add_paragraph(g.get('description', ''), style='List Bullet')
                p.runs[0].font.size = Pt(10)

        citations = data.get('citations', [])
        if citations:
            doc.add_heading(f'Citations ({len(citations)})', level=1)
            for c in citations:
                text = c.get('custom_format') or c.get('bibtex') or ''
                if text:
                    para = doc.add_paragraph(text)
                    para.runs[0].font.name = 'Courier New'
                    para.runs[0].font.size = Pt(8)

        doc.save(tmp_path)

        return send_file(
            tmp_path,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )

    except ImportError:
        return jsonify({'error': 'python-docx is not installed. Run: pip install python-docx'}), 500
    except Exception as e:
        app.logger.error(f'DOCX export error: {e}')
        import traceback; traceback.print_exc()
        return jsonify({'error': f'Word export failed: {str(e)}'}), 500


# ─── Plagiarism / Similarity Checker ───────────────────────────────────────

@app.route('/plagiarism')
def plagiarism_page():
    return render_template('plagiarism.html')


@limiter.limit("10 per minute")
@app.route('/check-plagiarism', methods=['POST'])
def check_plagiarism():
    try:
        data = request.get_json(force=True)
        text = _safe_text(data.get('text'))
        if not text:
            return jsonify({'error': 'No text provided'}), 400
        if len(text.split()) < 30:
            return jsonify({'error': 'Please provide at least 30 words for meaningful analysis'}), 400

        agent = PlagiarismCheckerAgent()
        result = agent.check(text)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f'Plagiarism check error: {e}')
        return jsonify({'error': str(e)}), 500


# ─── Journal Fit Recommender ────────────────────────────────────────────────

@app.route('/journal-fit')
def journal_fit_page():
    return render_template('journal_fit.html')


@limiter.limit("10 per minute")
@app.route('/recommend-journals', methods=['POST'])
def recommend_journals():
    try:
        data = request.get_json(force=True)
        abstract = _safe_text(data.get('abstract'))
        field = (data.get('field') or '').strip()
        if not abstract:
            return jsonify({'error': 'Abstract is required'}), 400

        agent = JournalRecommenderAgent()
        result = agent.recommend(abstract, field)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f'Journal recommendation error: {e}')
        return jsonify({'error': str(e)}), 500


# ─── Abstract Screener ──────────────────────────────────────────────────────

@app.route('/abstract-screener')
def abstract_screener_page():
    from src.agents.llm_client import get_provider
    has_key = get_provider() is not None
    return render_template('abstract_screener.html', has_anthropic_key=has_key)


@limiter.limit("10 per minute")
@app.route('/screen-abstracts', methods=['POST'])
def screen_abstracts():
    try:
        data = request.get_json(force=True)
        abstracts = data.get('abstracts') or []
        criteria = _safe_text(data.get('criteria'))
        if not abstracts:
            return jsonify({'error': 'No abstracts provided'}), 400
        if not criteria:
            return jsonify({'error': 'Inclusion/exclusion criteria are required'}), 400
        if len(abstracts) > 100:
            return jsonify({'error': 'Maximum 100 abstracts per request'}), 400

        agent = AbstractScreenerAgent()
        result = agent.screen(abstracts, criteria)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f'Abstract screening error: {e}')
        return jsonify({'error': str(e)}), 500


# ─── Pre-Submission Checklist ───────────────────────────────────────────────

@app.route('/checklist')
def checklist_page():
    from src.agents.llm_client import get_provider
    has_key = get_provider() is not None
    return render_template('checklist.html', has_anthropic_key=has_key)


@limiter.limit("10 per minute")
@app.route('/generate-checklist', methods=['POST'])
def generate_checklist():
    try:
        data = request.get_json(force=True)
        manuscript_type = (data.get('manuscript_type') or 'Research Article').strip()
        journal = (data.get('journal') or '').strip()
        field = (data.get('field') or '').strip()
        abstract = _safe_text(data.get('abstract'))

        agent = ChecklistAgent()
        result = agent.generate(manuscript_type, journal, field, abstract)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f'Checklist generation error: {e}')
        return jsonify({'error': str(e)}), 500


# ─── Grant Writer ───────────────────────────────────────────────────────────

@app.route('/grant-writer')
def grant_writer_page():
    from src.agents.llm_client import get_provider
    funders = GrantWriterAgent.get_funders()
    has_key = get_provider() is not None
    return render_template('grant_writer.html', funders=funders, has_anthropic_key=has_key)


@limiter.limit("10 per minute")
@app.route('/write-grant', methods=['POST'])
def write_grant():
    try:
        data = request.get_json(force=True)
        project_title = (data.get('project_title') or '').strip()
        description = (data.get('description') or '').strip()
        funder = (data.get('funder') or 'NSF').strip()
        field = (data.get('field') or '').strip()
        budget = (data.get('budget') or '').strip()
        duration = (data.get('duration') or '').strip()
        team = (data.get('team') or '').strip()

        if not project_title or not description:
            return jsonify({'error': 'Project title and description are required'}), 400

        agent = GrantWriterAgent()
        result = agent.write(project_title, description, funder, field, budget, duration, team)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f'Grant writing error: {e}')
        return jsonify({'error': str(e)}), 500


@app.route('/funders')
def get_funders():
    return jsonify(GrantWriterAgent.get_funders())


# ─── Citation Network Explorer ──────────────────────────────────────────────

@app.route('/citation-network')
def citation_network_page():
    return render_template('citation_network.html')


@limiter.limit("10 per minute")
@app.route('/get-citation-network', methods=['POST'])
def get_citation_network():
    try:
        data = request.get_json(force=True)
        query = (data.get('query') or '').strip()
        depth = int(data.get('depth', 1))
        if not query:
            return jsonify({'error': 'Paper title or ID is required'}), 400

        agent = CitationNetworkAgent()
        result = agent.get_network(query, depth)
        return jsonify(result)
    except Exception as e:
        app.logger.error(f'Citation network error: {e}')
        return jsonify({'error': str(e)}), 500


# ─── Benchmark Dashboard ─────────────────────────────────────────────────────
# Demonstrates the paper's core claim in the running app: the dedicated,
# separately-measured contradiction/gap-detection pipeline vs a naive
# single-LLM baseline, on a frozen worked-example corpus from eval/topics/.

def _eval_topics_dir() -> Path:
    return Path(__file__).resolve().parent / 'eval' / 'topics'


def _list_benchmark_topics():
    topics = []
    topics_dir = _eval_topics_dir()
    if topics_dir.exists():
        for sub in sorted(topics_dir.iterdir()):
            if sub.is_dir() and (sub / 'papers.json').exists():
                info = {}
                topic_json = sub / 'topic.json'
                if topic_json.exists():
                    try:
                        info = json.loads(topic_json.read_text())
                    except (OSError, json.JSONDecodeError):
                        info = {}
                topics.append({'slug': sub.name, 'topic': info.get('topic', sub.name)})
    return topics


@app.route('/benchmark')
def benchmark_page():
    return render_template('benchmark.html', topics=_list_benchmark_topics())


@limiter.limit("5 per minute")
@app.route('/run-benchmark', methods=['POST'])
def run_benchmark():
    try:
        data = request.get_json(force=True)
        slug = (data.get('slug') or '').strip()
        topics_by_slug = {t['slug']: t for t in _list_benchmark_topics()}
        if slug not in topics_by_slug:
            return jsonify({'error': 'Unknown worked-example topic. Run eval.snapshot_topic first.'}), 400

        topic_dir = _eval_topics_dir() / slug
        papers_payload = json.loads((topic_dir / 'papers.json').read_text())
        topic_text = topics_by_slug[slug]['topic']

        from src.models.data_models import PaperMetadata
        from src.agents import llm_client

        def run_pipeline():
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                frozen = [PaperMetadata(**{k: v for k, v in record.items() if k != 'paper_id'})
                          for record in papers_payload]
                system = AutonomousResearchSystem()
                results = loop.run_until_complete(system.research(topic_text, papers=frozen))
                return {
                    'claims': len(results.claims),
                    'contradictions': len(results.contradictions),
                    'gaps': len(results.research_gaps),
                    'usage': results.usage,
                    'run_metadata': results.run_metadata,
                    'sample_gaps': [g.description for g in results.research_gaps[:3]],
                }
            finally:
                loop.close()

        def run_naive_baseline():
            from eval.run_baseline import build_prompt, parse_response, to_predictions
            tracker = llm_client.UsageTracker()
            prompt = build_prompt(topic_text, papers_payload)
            try:
                raw = llm_client.chat(prompt, max_tokens=4000, temperature=0.0,
                                       tracker=tracker, purpose='baseline')
                parsed = parse_response(raw)
            except Exception:
                tracker.record_parse_failure('baseline')
                parsed = {'claims': [], 'contradictions': [], 'gaps': []}
            preds = to_predictions(parsed, papers_payload, tracker.summary())
            return {
                'claims': len(preds['claims']),
                'contradictions': len(preds['contradictions']),
                'gaps': len(preds['gaps']),
                'usage': preds['usage'],
                'sample_gaps': [g.get('description', '') for g in preds['gaps'][:3]],
            }

        return jsonify({
            'topic': topic_text,
            'slug': slug,
            'papers_analyzed': len(papers_payload),
            'dedicated_pipeline': run_pipeline(),
            'naive_baseline': run_naive_baseline(),
        })
    except Exception as e:
        app.logger.error(f'Benchmark run error: {e}')
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    # Create output directory
    Path("output").mkdir(exist_ok=True)
    
    print("🌐 Starting Research System Web Interface (Fixed Version)")
    print("📍 Open your browser to: http://localhost:5000")
    print("🧪 Test endpoint: http://localhost:5000/test")
    print("📖 Citations page: http://localhost:5000/citations")
    
    app.run(debug=True, host='0.0.0.0', port=5000, threaded=True)